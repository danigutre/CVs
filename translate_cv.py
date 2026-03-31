"""
Translates a .docx CV to Spanish using AWS Bedrock (Claude Opus 4.6)
while preserving the exact document formatting.

Strategy: Extract all text segments with unique IDs, send the entire document
in a single inference call for full context, then map translations back.
"""

import json
import os
import sys
from pathlib import Path

import boto3
from docx import Document


MODEL_ID = "us.anthropic.claude-opus-4-6-v1"


def get_bedrock_client():
    """Create a Bedrock Runtime client using AWS credentials."""
    return boto3.client(
        "bedrock-runtime",
        region_name=os.environ.get("AWS_REGION", "us-east-1"),
    )


# ---------------------------------------------------------------------------
# Phase 1: Extract every text segment from the document with a unique key
# ---------------------------------------------------------------------------

def extract_segments(doc):
    """
    Walk the entire document and collect every non-empty text run.
    Returns a list of dicts:
        { "id": "<unique key>", "text": "<original text>" }
    and a parallel list of (location_type, reference, run_index) so we can
    write translations back.
    """
    segments = []
    locations = []  # mirrors segments – how to write back

    seg_id = 0

    # --- body paragraphs ---
    for p_idx, paragraph in enumerate(doc.paragraphs):
        for r_idx, run in enumerate(paragraph.runs):
            if run.text.strip():
                segments.append({"id": f"B{seg_id}", "text": run.text})
                locations.append(("body", p_idx, r_idx))
                seg_id += 1

    # --- tables (may be nested) ---
    def _walk_table(table, table_path):
        nonlocal seg_id
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for p_idx, paragraph in enumerate(cell.paragraphs):
                    for r_idx, run in enumerate(paragraph.runs):
                        if run.text.strip():
                            segments.append({"id": f"B{seg_id}", "text": run.text})
                            locations.append(("table", (*table_path, row_idx, col_idx, p_idx), r_idx))
                            seg_id += 1
                for nt_idx, nested in enumerate(cell.tables):
                    _walk_table(nested, (*table_path, row_idx, col_idx, "nt", nt_idx))

    for t_idx, table in enumerate(doc.tables):
        _walk_table(table, (t_idx,))

    # --- headers / footers ---
    for s_idx, section in enumerate(doc.sections):
        for hf_attr in ("header", "footer", "first_page_header", "first_page_footer",
                         "even_page_header", "even_page_footer"):
            hf = getattr(section, hf_attr, None)
            if hf is None or hf.is_linked_to_previous:
                continue
            for p_idx, paragraph in enumerate(hf.paragraphs):
                for r_idx, run in enumerate(paragraph.runs):
                    if run.text.strip():
                        segments.append({"id": f"B{seg_id}", "text": run.text})
                        locations.append(("hf", (s_idx, hf_attr, p_idx), r_idx))
                        seg_id += 1
            for t_idx, table in enumerate(hf.tables):
                _walk_table(table, ("hf", s_idx, hf_attr, t_idx))

    return segments, locations


# ---------------------------------------------------------------------------
# Phase 2: Single-inference translation
# ---------------------------------------------------------------------------

def translate_all(client, segments):
    """
    Send every segment to the model in one call.
    Returns a dict  { segment_id: translated_text }.
    """
    if not segments:
        return {}

    # Build a compact JSON payload the model must mirror back
    payload = [{"id": s["id"], "text": s["text"]} for s in segments]

    prompt = (
        "You are translating a professional CV / résumé to Spanish.\n"
        "Below is a JSON array of text segments extracted from the document. "
        "Each segment has an \"id\" and \"text\".\n\n"
        "Translate every segment's \"text\" to Spanish. "
        "If a segment is already in Spanish, keep it unchanged.\n\n"
        "CRITICAL RULES:\n"
        "1. Return a JSON array with the EXACT same number of objects, in the EXACT same order.\n"
        "2. Each object must have \"id\" (unchanged) and \"text\" (translated).\n"
        "3. Preserve leading/trailing whitespace in each segment.\n"
        "4. Do NOT merge, split, add, or remove segments.\n"
        "5. Return ONLY the JSON array — no markdown fences, no explanation.\n\n"
        f"{json.dumps(payload, ensure_ascii=False)}"
    )

    print(f"Sending {len(segments)} segments in a single inference call...")

    response = client.invoke_model(
        modelId=MODEL_ID,
        contentType="application/json",
        accept="application/json",
        body=json.dumps({
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": 16384,
            "messages": [{"role": "user", "content": prompt}],
        }),
    )

    result = json.loads(response["body"].read())
    raw = result["content"][0]["text"].strip()

    # Parse the JSON array
    try:
        translated = json.loads(raw)
    except json.JSONDecodeError:
        # Try extracting the JSON array from surrounding text
        start = raw.find("[")
        end = raw.rfind("]") + 1
        if start != -1 and end > start:
            translated = json.loads(raw[start:end])
        else:
            print("ERROR: Could not parse model response as JSON.")
            print("Raw response (first 500 chars):", raw[:500])
            sys.exit(1)

    # Build lookup by id
    translation_map = {}
    for item in translated:
        translation_map[item["id"]] = item["text"]

    return translation_map


# ---------------------------------------------------------------------------
# Phase 3: Write translations back into the document
# ---------------------------------------------------------------------------

def _resolve_table_cell(doc, path):
    """Navigate nested table paths to reach the target cell/paragraph."""
    # path is a tuple like (table_idx, row, col, para) or with nested ("nt", idx, ...)
    obj = doc
    i = 0

    # Handle header/footer prefix
    if path[0] == "hf":
        section = doc.sections[path[1]]
        hf = getattr(section, path[2])
        # Rest of path starting from table index
        obj = hf
        i = 3

    while i < len(path):
        if path[i] == "nt":
            # nested table within current cell
            obj = obj.tables[path[i + 1]]
            i += 2
        elif isinstance(path[i], int):
            # table_idx -> row -> col -> para
            remaining = path[i:]
            table = obj.tables[remaining[0]]
            row = table.rows[remaining[1]]
            cell = row.cells[remaining[2]]
            para = cell.paragraphs[remaining[3]]
            return para
            break
        i += 1
    return None


def apply_translations(doc, segments, locations, translation_map):
    """Write translated text back into the document runs."""
    applied = 0
    skipped = 0

    for seg, loc in zip(segments, locations):
        seg_id = seg["id"]
        if seg_id not in translation_map:
            skipped += 1
            continue

        translated = translation_map[seg_id]
        loc_type = loc[0]

        try:
            if loc_type == "body":
                _, p_idx, r_idx = loc
                doc.paragraphs[p_idx].runs[r_idx].text = translated

            elif loc_type == "table":
                _, cell_path, r_idx = loc
                # cell_path is a tuple: (table_idx, row, col, para) possibly with nested
                # Navigate to the right place
                para = _navigate_to_para(doc, cell_path)
                if para:
                    para.runs[r_idx].text = translated

            elif loc_type == "hf":
                _, hf_path, r_idx = loc
                s_idx, hf_attr, p_idx = hf_path
                hf = getattr(doc.sections[s_idx], hf_attr)
                hf.paragraphs[p_idx].runs[r_idx].text = translated

            applied += 1
        except (IndexError, AttributeError) as e:
            print(f"  Warning: Could not apply segment {seg_id}: {e}")
            skipped += 1

    print(f"  Applied {applied} translations, {skipped} skipped.")


def _navigate_to_para(doc, cell_path):
    """Navigate a table cell path tuple to reach the target paragraph."""
    # cell_path examples:
    #   (table_idx, row, col, para_idx)
    #   (table_idx, row, col, "nt", nested_table_idx, row, col, para_idx)
    #   ("hf", section_idx, hf_attr, table_idx, row, col, para_idx)
    path = list(cell_path)
    i = 0

    # Determine root container (doc or header/footer)
    if path[0] == "hf":
        section = doc.sections[path[1]]
        container = getattr(section, path[2])
        i = 3
    else:
        container = doc
        i = 0

    # Navigate: table -> row -> cell, possibly nested
    table = container.tables[path[i]]
    i += 1

    while i < len(path):
        if path[i] == "nt":
            # We're in a cell already, go to nested table
            i += 1
            table = table.tables[path[i]] if hasattr(table, 'tables') else None
            if table is None:
                return None
            i += 1
        else:
            row_idx = path[i]
            col_idx = path[i + 1]
            para_idx = path[i + 2]

            cell = table.rows[row_idx].cells[col_idx]

            # Check if there's more nesting after para_idx
            if i + 3 < len(path) and path[i + 3] == "nt":
                table = cell.tables[path[i + 4]]
                i = i + 5
            else:
                return cell.paragraphs[para_idx]

    return None


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def translate_document(input_path: str, output_path: str = None):
    """Main function to translate a .docx document to Spanish."""
    input_path = Path(input_path)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}")
        sys.exit(1)

    if output_path is None:
        output_path = input_path.parent / f"{input_path.stem}_ES{input_path.suffix}"
    else:
        output_path = Path(output_path)

    print(f"Input:  {input_path}")
    print(f"Output: {output_path}")
    print()

    # Load document
    doc = Document(str(input_path))
    client = get_bedrock_client()

    # Phase 1: Extract all text segments
    print("Phase 1: Extracting text segments...")
    segments, locations = extract_segments(doc)
    print(f"  Found {len(segments)} text segments.")
    print()

    if not segments:
        print("No translatable text found.")
        return

    # Phase 2: Translate everything in one call
    print("Phase 2: Translating (single inference with full context)...")
    translation_map = translate_all(client, segments)
    print(f"  Received {len(translation_map)} translations.")
    print()

    # Phase 3: Write translations back
    print("Phase 3: Applying translations to document...")
    apply_translations(doc, segments, locations, translation_map)

    # Save
    doc.save(str(output_path))
    print(f"\nDone! Translated document saved to: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        script_dir = Path(__file__).parent
        docx_files = list(script_dir.glob("*.docx"))
        if not docx_files:
            print("Usage: python translate_cv.py <input.docx> [output.docx]")
            sys.exit(1)

        if len(docx_files) == 1:
            input_file = docx_files[0]
        else:
            print("Available .docx files:")
            for i, f in enumerate(docx_files, 1):
                print(f"  {i}. {f.name}")
            choice = int(input("Select file number: ")) - 1
            input_file = docx_files[choice]

        translate_document(str(input_file))
    else:
        input_file = sys.argv[1]
        output_file = sys.argv[2] if len(sys.argv) > 2 else None
        translate_document(input_file, output_file)
